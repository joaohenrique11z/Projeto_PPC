"""
services/diagrama_service.py

Gera DOIS diagramas de matriz curricular e os exporta em um único arquivo .docx:

  Diagrama 1 — Grade de disciplinas (HTML → PNG via Playwright)
    Cards coloridos por núcleo agrupados em colunas de período, com
    soma de carga horária no rodapé de cada coluna.

  Diagrama 2 — Fluxograma de pré-requisitos (Graphviz → PNG)
    Grafo direcionado com subgrafos por período, arestas ortogonais
    e cores por núcleo curricular.

Fluxo principal (gerar_matrizes_unificadas_docx):
  1. Buscar componentes e dependências do ppc_id no Supabase.
  2. Renderizar HTML da grade e capturar screenshot → matriz_grade.png.
  3. Montar grafo Graphviz e renderizar → matriz_fluxo.png.
  4. Criar um único .docx com ambas as imagens.
  5. Retornar caminhos dos arquivos temporários para limpeza posterior.
"""

import os
import tempfile
import textwrap

import graphviz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from database import supabase

# Garante que o executável dot do Graphviz seja encontrado no Windows.
_GRAPHVIZ_BIN = r"C:\Program Files\Graphviz\bin"
if os.path.isdir(_GRAPHVIZ_BIN) and _GRAPHVIZ_BIN not in os.environ.get("PATH", ""):
    os.environ["PATH"] = _GRAPHVIZ_BIN + os.pathsep + os.environ.get("PATH", "")


# ── Paleta de cores (compartilhada por ambos os diagramas) ───────────────────
CORES_NUCLEO: dict[str, str] = {
    "Básico":       "#AED6F1",   # azul claro
    "Específico":   "#A9DFBF",   # verde claro
    "Profissional": "#F9E79F",   # amarelo claro
    "Optativo":     "#F5CBA7",   # laranja claro
    "Extensão":     "#D7BDE2",   # lilás claro
}
COR_PADRAO = "#E8E8E8"  # cinza claro para núcleos não mapeados


# ── Helpers de banco de dados ─────────────────────────────────────────────────

def _buscar_ppc(ppc_id: str) -> dict:
    """
    Retorna os dados básicos do PPC (nome_curso, etc.).

    Evita .single() porque no supabase-py v2 esse método lança APIError
    quando nenhuma linha é encontrada, em vez de retornar lista vazia.
    """
    response = (
        supabase
        .table("ppc")
        .select("id, nome_curso")
        .eq("id", ppc_id)
        .limit(1)
        .execute()
    )
    rows = response.data or []
    return rows[0] if rows else {}


def _buscar_componentes(ppc_id: str) -> list[dict]:
    """
    Retorna todos os componentes curriculares de um PPC ordenados por período.

    Seleciona apenas colunas garantidas pelo schema base. O campo 'codigo'
    é buscado separadamente porque pode não existir em versões antigas do banco.

    Args:
        ppc_id: UUID do PPC.

    Returns:
        Lista de dicionários com os campos do componente_curricular.

    Raises:
        RuntimeError: Se a query falhar no Supabase (ex: coluna inexistente).
    """
    response = (
        supabase
        .table("componente_curricular")
        .select("id, codigo, nome, periodo, nucleo_curricular, ch_total_relogio")
        .eq("ppc_id", ppc_id)
        .order("periodo")
        .execute()
    )

    # response.data é None (não lista vazia) quando a query falha silenciosamente.
    # Isso pode ocorrer se uma coluna solicitada não existir no banco.
    if response.data is None:
        # Tenta query mínima para verificar se o problema é a coluna 'codigo'
        fallback = (
            supabase
            .table("componente_curricular")
            .select("id, nome, periodo, nucleo_curricular, ch_total_relogio")
            .eq("ppc_id", ppc_id)
            .order("periodo")
            .execute()
        )
        if fallback.data is None:
            raise RuntimeError(
                f"Erro ao buscar componentes do PPC '{ppc_id}'. "
                "Verifique as colunas da tabela componente_curricular."
            )
        # Normaliza: adiciona 'codigo' vazio para componentes sem essa coluna
        for comp in (fallback.data or []):
            comp.setdefault("codigo", "")
        return fallback.data or []

    return response.data or []


def _buscar_dependencias(ppc_id: str) -> list[dict]:
    """
    Retorna todas as dependências entre componentes que pertencem ao PPC.

    Usa duas queries em vez de join com nome de FK explícito, pois o nome
    da constraint pode variar entre ambientes e causar erros 400 no Supabase.

    Args:
        ppc_id: UUID do PPC.

    Returns:
        Lista de dicts com componente_base_id e componente_alvo_id.
    """
    # Passo 1: coleta os IDs de todos os componentes deste PPC
    comp_response = (
        supabase
        .table("componente_curricular")
        .select("id")
        .eq("ppc_id", ppc_id)
        .execute()
    )
    ids_do_ppc = {
        row["id"] for row in (comp_response.data or [])
    }

    if not ids_do_ppc:
        return []

    # Passo 2: busca dependências onde o componente_base pertence a este PPC.
    # Usa .in_() com a lista de IDs para evitar joins com nome de FK.
    dep_response = (
        supabase
        .table("componente_dependencia")
        .select("componente_base_id, componente_alvo_id")
        .in_("componente_base_id", list(ids_do_ppc))
        .execute()
    )
    return dep_response.data or []


def _agrupados_por_periodo(componentes: list[dict]) -> dict[int, list[dict]]:
    """
    Organiza os componentes em um dicionário indexado pelo período.

    Args:
        componentes: Lista de componentes curriculares.

    Returns:
        Dict { periodo: [componente, ...] }.
    """
    grupos: dict[int, list[dict]] = {}
    for comp in componentes:
        periodo = comp.get("periodo", 0)
        grupos.setdefault(periodo, []).append(comp)
    return grupos


# ── Diagrama 1: Grade HTML → PNG ──────────────────────────────────────────────

def _prerequisitos_por_alvo(dependencias: list[dict]) -> dict[str, list[str]]:
    """
    Constrói um mapa de componente_alvo_id → [componente_base_id, ...].

    Usado para exibir as siglas dos pré-requisitos no card da grade.

    Args:
        dependencias: Lista de dependências retornada por _buscar_dependencias.

    Returns:
        Dict { alvo_id: [base_id, ...] }.
    """
    mapa: dict[str, list[str]] = {}
    for dep in dependencias:
        alvo = dep["componente_alvo_id"]
        base = dep["componente_base_id"]
        mapa.setdefault(alvo, []).append(base)
    return mapa


def _renderizar_html_grade(
    nome_curso: str,
    grupos: dict[int, list[dict]],
    prereqs_por_alvo: dict[str, list[str]],
    id_para_codigo: dict[str, str],
) -> str:
    """
    Gera o HTML completo da grade curricular usando CSS Grid.

    Cada coluna representa um período; cada célula é um card colorido
    com código, pré-requisitos, nome e carga horária.

    Args:
        nome_curso: Nome do curso para o título.
        grupos: Dict { periodo: [componente, ...] }.
        prereqs_por_alvo: Dict { alvo_id: [base_id, ...] }.
        id_para_codigo: Dict { id: codigo } para resolver os pré-requisitos.

    Returns:
        String com o HTML completo pronto para screenshot.
    """
    periodos_ordenados = sorted(grupos.keys())
    num_periodos = len(periodos_ordenados)

    # Cabeçalhos de período
    cabecalhos_html = ""
    for p in periodos_ordenados:
        cabecalhos_html += f'<div class="header-cell">{p}º Período</div>\n'

    # Cards por coluna
    colunas_html = ""
    for p in periodos_ordenados:
        colunas_html += '<div class="column">\n'
        total_ch = 0

        for comp in grupos[p]:
            cor = CORES_NUCLEO.get(comp.get("nucleo_curricular") or "", COR_PADRAO)
            ch = comp.get("ch_total_relogio") or 0
            total_ch += ch
            codigo = comp.get("codigo") or ""
            nome = comp.get("nome") or "—"

            # Siglas dos pré-requisitos
            prereq_ids = prereqs_por_alvo.get(comp["id"], [])
            prereq_codigos = [id_para_codigo.get(pid, "?") for pid in prereq_ids]
            prereq_html = (
                f'<span class="prereq">{", ".join(prereq_codigos)}</span>'
                if prereq_codigos else ""
            )

            colunas_html += textwrap.dedent(f"""\
                <div class="card" style="background:{cor}">
                  <div class="card-top">
                    <span class="codigo">{codigo}</span>
                    {prereq_html}
                  </div>
                  <div class="card-nome">{nome}</div>
                  <div class="card-ch">{ch} h/r</div>
                </div>
            """)

        colunas_html += f'<div class="col-total">{total_ch} h/r</div>\n'
        colunas_html += "</div>\n"

    html = textwrap.dedent(f"""\
        <!DOCTYPE html>
        <html lang="pt-BR">
        <head>
          <meta charset="UTF-8"/>
          <style>
            * {{ box-sizing: border-box; margin: 0; padding: 0; }}
            body {{
              font-family: Arial, sans-serif;
              background: #f5f5f5;
              padding: 24px;
            }}
            h1 {{
              text-align: center;
              font-size: 16px;
              margin-bottom: 20px;
              color: #1a1a2e;
              text-transform: uppercase;
              letter-spacing: 1px;
            }}
            .grid {{
              display: grid;
              grid-template-columns: repeat({num_periodos}, 1fr);
              gap: 6px;
              width: 100%;
            }}
            .header-cell {{
              background: #1a1a2e;
              color: #fff;
              text-align: center;
              padding: 8px 4px;
              font-size: 11px;
              font-weight: bold;
              border-radius: 4px;
            }}
            .column {{
              display: flex;
              flex-direction: column;
              gap: 5px;
            }}
            .card {{
              border-radius: 5px;
              padding: 7px 8px;
              border: 1px solid rgba(0,0,0,0.12);
              display: flex;
              flex-direction: column;
              gap: 4px;
              min-height: 80px;
            }}
            .card-top {{
              display: flex;
              justify-content: space-between;
              align-items: flex-start;
            }}
            .codigo {{
              font-size: 9px;
              font-weight: bold;
              color: #333;
            }}
            .prereq {{
              font-size: 8px;
              color: #c0392b;
              font-style: italic;
              text-align: right;
              max-width: 60%;
            }}
            .card-nome {{
              font-size: 10px;
              font-weight: bold;
              color: #1a1a2e;
              flex: 1;
              line-height: 1.3;
            }}
            .card-ch {{
              font-size: 9px;
              color: #555;
              text-align: right;
            }}
            .col-total {{
              background: #1a1a2e;
              color: #fff;
              text-align: center;
              padding: 5px;
              font-size: 10px;
              font-weight: bold;
              border-radius: 4px;
            }}
          </style>
        </head>
        <body>
          <h1>Matriz Curricular — {nome_curso}</h1>
          <div class="grid">
            {cabecalhos_html}
            {colunas_html}
          </div>
        </body>
        </html>
    """)
    return html


def _screenshot_html_para_png(html_content: str, output_path: str) -> None:
    """
    Usa Playwright para renderizar o HTML e salvar um screenshot PNG.

    Playwright roda headless (sem janela visível) e suporta páginas
    com layouts complexos de CSS Grid/Flexbox.

    Args:
        html_content: String com o HTML completo.
        output_path: Caminho absoluto onde o PNG será salvo.

    Raises:
        RuntimeError: Se o Playwright não estiver instalado corretamente.
    """
    try:
        from playwright.sync_api import sync_playwright
    except ImportError as err:
        raise RuntimeError(
            "Playwright não está instalado. Execute: "
            "pip install playwright && playwright install chromium"
        ) from err

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1600, "height": 900})
        page.set_content(html_content, wait_until="networkidle")
        # Ajusta altura ao conteúdo real da página
        page.evaluate("document.body.style.height = 'auto'")
        page.screenshot(path=output_path, full_page=True)
        browser.close()


# ── Diagrama 2: Graphviz → PNG ────────────────────────────────────────────────

def _label_no_graphviz(comp: dict) -> str:
    """
    Monta o label do nó Graphviz com nome e carga horária.

    Args:
        comp: Dicionário com dados do componente.

    Returns:
        String de label formatada.
    """
    nome = comp.get("nome", "?")
    codigo = comp.get("codigo", "")
    ch = comp.get("ch_total_relogio") or 0
    return f"{codigo}\\n{nome}\\n{ch} h/r"


def _gerar_png_graphviz(
    componentes: list[dict],
    dependencias: list[dict],
    output_path: str,
) -> None:
    """
    Gera o fluxograma de pré-requisitos com Graphviz e salva como PNG.

    Subgrafos por período garantem o alinhamento horizontal de cada
    semestre; arestas ortogonais (splines=ortho) mantêm o layout limpo.

    Args:
        componentes: Lista de componentes curriculares.
        dependencias: Lista de dependências entre componentes.
        output_path: Caminho base (sem extensão) onde o PNG será salvo.
                     O Graphviz adiciona ".png" automaticamente.
    """
    grafo = graphviz.Digraph(
        name="matriz_fluxo",
        graph_attr={
            "rankdir":  "TB",
            "splines":  "ortho",
            "nodesep":  "0.8",
            "ranksep":  "1.0",
            "fontname": "Helvetica",
            "bgcolor":  "white",
        },
        node_attr={
            "shape":    "box",
            "style":    "filled,rounded",
            "fontname": "Helvetica",
            "fontsize": "9",
            "margin":   "0.2",
        },
        edge_attr={
            "arrowsize": "0.7",
            "color":     "#555555",
        },
    )

    grupos = _agrupados_por_periodo(componentes)

    for periodo in sorted(grupos.keys()):
        with grafo.subgraph(name=f"cluster_periodo_{periodo}") as sub:
            sub.attr(
                label=f"  {periodo}º Período  ",
                style="rounded,filled",
                color="#CCCCCC",
                fillcolor="#F7F7F7",
                fontname="Helvetica",
                fontsize="10",
                fontcolor="#333333",
            )
            sub.attr(rank="same")
            for comp in grupos[periodo]:
                nucleo = comp.get("nucleo_curricular") or ""
                cor = CORES_NUCLEO.get(nucleo, COR_PADRAO)
                sub.node(
                    comp["id"],
                    label=_label_no_graphviz(comp),
                    fillcolor=cor,
                )

    for dep in dependencias:
        grafo.edge(dep["componente_base_id"], dep["componente_alvo_id"])

    grafo.render(
        filename=output_path,
        format="png",
        cleanup=True,   # remove o arquivo .gv intermediário
    )


# ── Helpers de formatação do Word ─────────────────────────────────────────────

def _adicionar_secao_paisagem(doc: Document) -> None:
    """
    Configura a última seção do documento para orientação paisagem.

    Usa manipulação direta de XML do OOXML porque python-docx não
    expõe essa propriedade nativamente na API pública.

    Args:
        doc: Instância de Document do python-docx.
    """
    section = doc.sections[-1]
    section.orientation = 1          # WD_ORIENT.LANDSCAPE = 1
    # Troca largura e altura
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height


def _adicionar_subtitulo(doc: Document, texto: str) -> None:
    """
    Adiciona um parágrafo de subtítulo (Heading level 2) ao documento.

    Args:
        doc: Instância de Document do python-docx.
        texto: Texto do subtítulo.
    """
    para = doc.add_heading(texto, level=2)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ── Função pública principal ──────────────────────────────────────────────────

def gerar_matrizes_unificadas_docx(ppc_id: str) -> tuple[str, list[str]]:
    """
    Orquestra a geração dos dois diagramas e os une em um único .docx.

    Etapas:
      1. Busca dados no Supabase.
      2. Gera Diagrama 1 (grade HTML → PNG via Playwright).
      3. Gera Diagrama 2 (fluxo Graphviz → PNG).
      4. Monta o documento Word com ambas as imagens.

    Args:
        ppc_id: UUID do PPC a ser processado.

    Returns:
        Tupla (caminho_docx, [caminho_png_grade, caminho_png_fluxo])
        com os arquivos temporários gerados.

    Raises:
        ValueError: Se nenhum componente for encontrado para o ppc_id.
        RuntimeError: Se o Playwright não estiver instalado.
    """
    ppc = _buscar_ppc(ppc_id)
    nome_curso = ppc.get("nome_curso") or "Curso sem nome"

    componentes = _buscar_componentes(ppc_id)
    if not componentes:
        raise ValueError(
            f"Nenhum componente curricular encontrado para o PPC '{ppc_id}'."
        )

    dependencias = _buscar_dependencias(ppc_id)
    grupos = _agrupados_por_periodo(componentes)

    # Mapa auxiliar id → codigo (para exibir pré-req nos cards)
    id_para_codigo: dict[str, str] = {
        c["id"]: (c.get("codigo") or c.get("nome", "?")[:6])
        for c in componentes
    }
    prereqs_por_alvo = _prerequisitos_por_alvo(dependencias)

    tmp_dir = tempfile.mkdtemp()

    # ── Diagrama 1: Grade HTML → PNG ─────────────────────────────────────────
    html_grade = _renderizar_html_grade(
        nome_curso, grupos, prereqs_por_alvo, id_para_codigo
    )
    png_grade_path = os.path.join(tmp_dir, "matriz_grade.png")
    _screenshot_html_para_png(html_grade, png_grade_path)

    # ── Diagrama 2: Graphviz → PNG ────────────────────────────────────────────
    png_fluxo_base = os.path.join(tmp_dir, "matriz_fluxo")
    _gerar_png_graphviz(componentes, dependencias, png_fluxo_base)
    png_fluxo_path = f"{png_fluxo_base}.png"

    # ── Montagem do documento Word ────────────────────────────────────────────
    doc = Document()

    # Título principal
    titulo = doc.add_heading(f"Matriz Curricular — {nome_curso}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # espaço

    # --- Seção 1: Grade de disciplinas (paisagem para caber melhor) ----------
    _adicionar_subtitulo(doc, "1. Visão em Grade (Estrutura Curricular)")
    doc.add_paragraph()

    # Insere a imagem ocupando toda a largura disponível em paisagem (~9 pol)
    doc.add_picture(png_grade_path, width=Inches(9))

    doc.add_page_break()

    # --- Seção 2: Fluxograma de dependências ---------------------------------
    _adicionar_subtitulo(doc, "2. Fluxo de Dependências e Pré-Requisitos")
    doc.add_paragraph()

    doc.add_picture(png_fluxo_path, width=Inches(7))

    # Legenda de núcleos
    doc.add_paragraph()
    legenda_titulo = doc.add_paragraph()
    legenda_titulo.add_run("Legenda de Núcleos Curriculares:").bold = True

    for nucleo, cor in CORES_NUCLEO.items():
        doc.add_paragraph(f"  ■  {nucleo}  (cor: {cor})")

    # Salva o arquivo
    docx_path = os.path.join(tmp_dir, "matrizes_curriculares.docx")
    doc.save(docx_path)

    arquivos_temporarios = [png_grade_path, png_fluxo_path]
    return docx_path, arquivos_temporarios


# ── Mantém a função original para não quebrar a rota existente ───────────────

def gerar_diagrama_docx(ppc_id: str) -> tuple[str, str]:
    """
    Compatibilidade: gera apenas o Diagrama 2 (Graphviz) em um .docx simples.

    Mantida para não quebrar a rota GET /api/exportar-diagrama/{ppc_id}.

    Args:
        ppc_id: UUID do PPC.

    Returns:
        Tupla (caminho_docx, caminho_png).
    """
    componentes = _buscar_componentes(ppc_id)
    if not componentes:
        raise ValueError(
            f"Nenhum componente curricular encontrado para o PPC '{ppc_id}'."
        )

    dependencias = _buscar_dependencias(ppc_id)
    tmp_dir = tempfile.mkdtemp()

    png_base = os.path.join(tmp_dir, "diagrama")
    _gerar_png_graphviz(componentes, dependencias, png_base)
    png_path = f"{png_base}.png"

    doc = Document()
    titulo = doc.add_heading("Matriz Curricular — Diagrama de Pré-Requisitos", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_picture(png_path, width=Inches(7))

    doc.add_paragraph()
    legenda = doc.add_paragraph()
    legenda.add_run("Legenda de Núcleos:").bold = True
    for nucleo, cor in CORES_NUCLEO.items():
        doc.add_paragraph(f"   ■ {nucleo}  (cor: {cor})")

    docx_path = os.path.join(tmp_dir, "matriz_curricular.docx")
    doc.save(docx_path)
    return docx_path, png_path
