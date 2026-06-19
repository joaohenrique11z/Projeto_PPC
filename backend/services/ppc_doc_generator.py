"""
services/ppc_doc_generator.py

Gerador de documentos DOCX para o PPC.

Fluxo:
  1. fetch_ppc_data()            → busca payload via carregar_ppc()
  2. build_context()             → converte payload em contexto plano para docxtpl
  3. render_template()           → aplica variáveis simples com docxtpl/Jinja2
  4. _replace_table_placeholders() → substitui tags {tabela_*} por tabelas reais
  5. generate_document()         → orquestra e retorna o caminho do arquivo gerado
"""

import logging
import os
import re
import tempfile
from collections import defaultdict
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docxtpl import DocxTemplate

from services.ppc_service import carregar_ppc

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Caminhos
# ─────────────────────────────────────────────────────────────────────────────

_BASE_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATE_PATH = _BASE_DIR / "templates" / "doc_ppc_modelo.docx"
EXPORTS_DIR = _BASE_DIR / "exports"

# Tags que representam tabelas geradas programaticamente (não são variáveis Jinja2)
_TABLE_TAGS: frozenset[str] = frozenset({
    "tabela_isntituicao",
    "tabela_curso",
    "tabela_indicadores_qualidade_curso",
    "tabela_divisao_componentes_por_nucleo",
    "tabela_componentes_por_periodo",
    "tabela_integralizacao",
    "tabela_optativas",
    "tabela_requisitos",
    "tabela_ementa",
    "tabela_coordenacao",
    "tabela_docentes",
    "tabela_composicao_colegiado",
    "tabela_composicao_nde",
    "fluxograma_periodos",
    "fluxograma_relacoes",
})


# ─────────────────────────────────────────────────────────────────────────────
# Exceção customizada
# ─────────────────────────────────────────────────────────────────────────────

class DocumentGenerationError(Exception):
    """Raised when document generation fails at any stage."""


# ─────────────────────────────────────────────────────────────────────────────
# 1. Busca de dados
# ─────────────────────────────────────────────────────────────────────────────

def fetch_ppc_data(ppc_id: str) -> dict:
    """
    Busca os dados completos do PPC diretamente do serviço de persistência.

    Args:
        ppc_id: UUID do PPC.

    Returns:
        Payload completo (ppc, coordenacao, membros, componentes, docentes, ambientes).

    Raises:
        DocumentGenerationError: Se o PPC não for encontrado ou ocorrer erro.
    """
    try:
        return carregar_ppc(ppc_id)
    except ValueError as exc:
        raise DocumentGenerationError(f"PPC {ppc_id} não encontrado: {exc}") from exc
    except Exception as exc:
        raise DocumentGenerationError(f"Erro ao buscar dados do PPC: {exc}") from exc


# ─────────────────────────────────────────────────────────────────────────────
# 2. Construção do contexto
# ─────────────────────────────────────────────────────────────────────────────

def _lookup_member(membros: list[dict], tipo: str, cargo: str) -> str:
    """Retorna o nome do primeiro membro que bate tipo e cargo (case-insensitive)."""
    for m in membros:
        if (m.get("tipo", "").lower() == tipo.lower()
                and m.get("cargo", "").lower() == cargo.lower()):
            return m.get("nome", "") or ""
    return ""


def _str(value: Any) -> str:
    """Converte valor para string, retornando '' para None."""
    return str(value) if value is not None else ""


def build_context(payload: dict) -> dict:
    """
    Converte o payload da API em um dicionário de contexto plano
    compatível com docxtpl/Jinja2.

    Todos os campos ausentes resultam em string vazia para evitar
    erros de renderização no template.

    Args:
        payload: Dicionário retornado por carregar_ppc().

    Returns:
        Dicionário com todas as variáveis de texto simples do template.
        As coleções (_ppc, _componentes, etc.) são incluídas para uso
        posterior pela camada de geração de tabelas.
    """
    ppc: dict = payload.get("ppc") or {}
    coordenacao: dict = payload.get("coordenacao") or {}
    membros: list[dict] = payload.get("membros") or []
    componentes: list[dict] = payload.get("componentes") or []
    docentes: list[dict] = payload.get("docentes") or []
    ambientes: list[dict] = payload.get("ambientes") or []

    # Extrai apenas o ano de data_ultima_atualizacao
    data_atualizacao = ppc.get("data_ultima_atualizacao", "")
    try:
        ano_atualizacao = str(
            datetime.fromisoformat(data_atualizacao.replace("Z", "+00:00")).year
        ) if data_atualizacao else ""
    except (ValueError, AttributeError):
        ano_atualizacao = ""

    # Converte semestres → anos (arredondado para 1 casa)
    min_sem = ppc.get("integralizacao_min_semestres") or 0
    max_sem = ppc.get("integralizacao_max_semestres") or 0

    def coleg(cargo: str) -> str:
        return _lookup_member(membros, "Colegiado", cargo)

    def com(cargo: str) -> str:
        return _lookup_member(membros, "Comissão de Elaboração", cargo)

    context: dict[str, Any] = {
        # ── Dados gerais do PPC ──────────────────────────────────────────────
        "nivel":                    _str(ppc.get("nivel")),
        "nome_curso":               _str(ppc.get("nome_curso")),
        "ano_atualizacao":          ano_atualizacao,
        "campus":                   _str(ppc.get("campus_name")),
        "tipo_curso":               _str(ppc.get("tipo_curso")),
        "eixo_tecnologico":         _str(ppc.get("eixo_tecnologico")),
        "area_conhecimento":        _str(ppc.get("area_conhecimento")),
        "inicio_curso":             _str(ppc.get("inicio_curso")),
        "regime_matricula":         _str(ppc.get("regime_matricula")),
        "ch_total_relogio":         _str(ppc.get("ch_total_relogio")),
        "ch_total_aula":            _str(ppc.get("ch_total_aula")),
        "duracao_aula_minutos":     _str(ppc.get("duracao_aula_minutos")),
        "semanas_letivas":          _str(ppc.get("semanas_letivas")),
        "integralizacao_min_anos":  _str(round(min_sem / 2, 1)) if min_sem else "",
        "integralizacao_max_anos":  _str(round(max_sem / 2, 1)) if max_sem else "",
        "x":                        _str(ppc.get("vagas_anuais")),    # vagas anuais
        "y":                        _str(ppc.get("vagas_turno")),     # vagas por turno
        "cursos_tecnicos_afins":    _str(ppc.get("cursos_tecnicos_afins")),
        "outros_cursos_campus":     _str(ppc.get("outros_cursos_campus")),
        "portaria":                 _str(ppc.get("portaria")),
        # Alias de segurança — cobre variações de tag no template
        "data_ultima_atualizacao":   ano_atualizacao,

        # ── Colegiado institucional ──────────────────────────────────────────
        "colegiado_reitor":                         coleg("Reitor"),
        "colegiado_pro_reitor_ensino":              coleg("Pró-Reitora de Ensino"),
        "colegiado_pro_reitor_pesquisa":            coleg("Pró-Reitora de Pesquisa, Pós-Graduação e Inovação"),
        "colegiado_pro_reitor_extensao":            coleg("Pró-Reitora de Extensão"),
        "colegiado_pro_reitor_integracao":          coleg("Pró-Reitora de Integração e Desenvolvimento Institucional"),
        "colegiado_pro_reitor_administracao":       coleg("Pró-Reitor de Administração"),
        "colegiado_diretor_geral":                  coleg("Diretor Geral do Campus"),
        "colegiado_diretor_adm":                    coleg("Diretora de Administração e Planejamento"),
        "colegiado_diretor_des":                    coleg("Diretor de Desenvolvimento Educacional"),
        "colegiado_gestao_acad":                    coleg("Departamento de Gestão Acadêmica e Produção"),
        "colegiado_coord_registro":                 coleg("Coordenador de Registro Acadêmico e Diplomação"),
        "colegiado_coord_extensao":                 coleg("Coordenador de Extensão"),
        "colegiado_coord_pesquisa":                 coleg("Coordenador de Pesquisa, Pós-Graduação e Inovação"),
        "colegiado_coord_curso":                    coleg("Coordenador do Curso"),
        "colegiado_assessoria_ped":                 coleg("Assessoria Pedagógica"),

        # ── Comissão de elaboração ───────────────────────────────────────────
        "presidente_comissao_de_elaboracao":        com("Presidente da Comissão de Elaboração do PPC"),
        "comissao_de_elaboracao_membro":            com("Membro da Comissão"),
        "comissao_de_elaboracao_bibliotecario":     com("Bibliotecária"),
        "comissao_de_elaboracao_pedagogo":          com("Pedagogo"),
        "comissao_de_elaboracao_revisao_textual":   com("Responsável pela Revisão Textual"),

        # ── Coleções brutas para geração de tabelas ──────────────────────────
        # (prefixo _ indica uso interno, não são variáveis de template)
        "_ppc":           ppc,
        "_coordenacao":   coordenacao,
        "_membros":       membros,
        "_componentes":   componentes,
        "_docentes":      docentes,
        "_ambientes":     ambientes,
    }

    logger.debug("Contexto construído com %d chaves de texto.", len([k for k in context if not k.startswith("_")]))
    return context


# ─────────────────────────────────────────────────────────────────────────────
# 3. Renderização via docxtpl (variáveis de texto simples)
# ─────────────────────────────────────────────────────────────────────────────

def _prepare_jinja_template(src: Path, dst: Path) -> None:
    """
    Converte o template original (chaves simples {var}) para o formato
    Jinja2 (chaves duplas {{ var }}), preservando as tags de tabelas.

    Esta conversão é feita em memória e salva no caminho destino.
    Como o Word pode quebrar variáveis entre runs, a estratégia é
    reconstruir cada parágrafo com o texto completo no primeiro run.

    Args:
        src: Caminho do template original (.docx).
        dst: Caminho de saída do template convertido.
    """
    doc = Document(str(src))

    # Mapeamentos especiais: tag original → variável Jinja2 de destino
    # Necessário para tags que contêm parênteses ou outros caracteres
    # que o Jinja2 interpretaria como sintaxe de função.
    _SPECIAL_MAPPINGS: dict[str, str] = {
        "{data_ultima_atualizacao(ano)}": "{{ ano_atualizacao }}",
        "{ regime_matricula}": "{{ regime_matricula }}",
    }

    def convert_para(para: Any) -> None:
        full_text = para.text
        if not full_text.strip():
            return

        # Verifica se é uma tag de tabela — mantém intacta
        stripped = full_text.strip()
        inner = stripped.lstrip("{").rstrip("}").strip()
        if inner in _TABLE_TAGS:
            return

        new_text = full_text

        # Aplica mapeamentos especiais antes da regex genérica
        for original, replacement in _SPECIAL_MAPPINGS.items():
            if original in new_text:
                new_text = new_text.replace(original, replacement)

        # Substitui {var} → {{ var }} nas variáveis de texto restantes.
        # Pontos nas chaves (ex: {colegiado.reitor}) são substituídos por underscore
        # para corresponder ao contexto plano gerado por build_context().
        # O charset inclui unicode (\w com flag UNICODE) para cobrir acentos como ã.
        def replace_var(m: re.Match) -> str:
            key = m.group(1).replace(".", "_")
            return f"{{{{ {key} }}}}"

        new_text = re.sub(
            r"(?<![{])\{\s*([\w][\w.]*?)\s*\}(?![}])",
            replace_var,
            new_text,
            flags=re.UNICODE,
        )

        if new_text == full_text:
            return

        # Consolida todo o texto no primeiro run, apaga os demais
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        convert_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    convert_para(para)

    doc.save(str(dst))
    logger.info("Template Jinja2 preparado em %s", dst)


def render_template(jinja_template_path: Path, context: dict, output_path: Path) -> None:
    """
    Aplica o contexto de variáveis simples no template usando docxtpl.

    As chaves internas (prefixo _) são removidas antes da renderização
    para evitar erros do Jinja2 com tipos não serializáveis.

    Args:
        jinja_template_path: Caminho do template já convertido para Jinja2.
        context:             Dicionário de contexto retornado por build_context().
        output_path:         Caminho onde o DOCX intermediário será salvo.
    """
    # Remove chaves internas — docxtpl não deve processá-las
    jinja_context = {k: v for k, v in context.items() if not k.startswith("_")}

    tpl = DocxTemplate(str(jinja_template_path))
    tpl.render(jinja_context, autoescape=False)
    tpl.save(str(output_path))
    logger.info("Template renderizado (variáveis de texto) → %s", output_path)


# ─────────────────────────────────────────────────────────────────────────────
# 4. Geração de tabelas via python-docx
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell: Any, hex_color: str) -> None:
    """Define cor de fundo de uma célula via XML."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _bold_cell(cell: Any, text: str, bg_hex: str | None = None) -> None:
    """Escreve texto em negrito em uma célula, com fundo opcional."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    if bg_hex:
        _set_cell_bg(cell, bg_hex)


def _merge_row(row: Any) -> None:
    """Mescla todas as células de uma linha horizontalmente."""
    if len(row.cells) > 1:
        row.cells[0].merge(row.cells[-1])


def _insert_table_after_paragraph(table: Any, paragraph: Any) -> None:
    """Insere o elemento XML da tabela imediatamente após o parágrafo."""
    paragraph._element.addnext(table._element)


def _remove_paragraph(paragraph: Any) -> None:
    """Remove um parágrafo do documento."""
    p_el = paragraph._element
    p_el.getparent().remove(p_el)


# ── Builders individuais de tabela ───────────────────────────────────────────

def _build_tabela_instituicao(doc: Any, ppc: dict) -> Any:
    """Tabela 2 colunas: Campo / Valor — dados institucionais do campus."""
    fields = [
        ("Instituição", "Instituto Federal de Educação, Ciência e Tecnologia de Pernambuco"),
        ("Campus", ppc.get("campus_name", "")),
        ("CNPJ", ppc.get("cnpj", "")),
        ("CEP", ppc.get("cep", "")),
        ("Cidade", ppc.get("cidade", "")),
        ("Bairro", ppc.get("bairro", "")),
        ("Endereço", f'{ppc.get("rua", "")}, {ppc.get("numero", "")}'.strip(", ")),
        ("Telefone/Fax", ppc.get("telefone_fax", "")),
        ("E-mail", ppc.get("email_contato", "")),
        ("Ato Legal de Criação", ppc.get("ato_legal", "")),
        ("Sítio Eletrônico", ppc.get("sitio_web", "")),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    for i, (campo, valor) in enumerate(fields):
        _bold_cell(table.rows[i].cells[0], campo, "D9D9D9")
        table.rows[i].cells[1].text = _str(valor)
    return table


def _build_tabela_curso(doc: Any, ppc: dict) -> Any:
    """Tabela 2 colunas: Campo / Valor — dados de identificação do curso."""
    min_sem = ppc.get("integralizacao_min_semestres") or 0
    max_sem = ppc.get("integralizacao_max_semestres") or 0
    fields = [
        ("Curso", ppc.get("nome_curso", "")),
        ("Nível", ppc.get("nivel", "")),
        ("Tipo de Curso", ppc.get("tipo_curso", "")),
        ("Modalidade", ppc.get("modalidade_curso", "")),
        ("Titulação", ppc.get("titulacao", "")),
        ("Área do Conhecimento", ppc.get("area_conhecimento", "")),
        ("Eixo Tecnológico", ppc.get("eixo_tecnologico", "")),
        ("Carga Horária Total (h/r)", _str(ppc.get("ch_total_relogio"))),
        ("Carga Horária Total (h/a)", _str(ppc.get("ch_total_aula"))),
        ("Duração da Aula (min)", _str(ppc.get("duracao_aula_minutos"))),
        ("CH Extensão (h/r)", _str(ppc.get("ch_extensao"))),
        ("Atividades Complementares (h/r)", _str(ppc.get("atividades_complementares"))),
        ("Integralização Mínima", f"{round(min_sem / 2, 1)} ano(s) / {min_sem} semestre(s)"),
        ("Integralização Máxima", f"{round(max_sem / 2, 1)} ano(s) / {max_sem} semestre(s)"),
        ("Semanas Letivas por Semestre", _str(ppc.get("semanas_letivas"))),
        ("Turno(s)", ppc.get("turnos", "")),
        ("Vagas Anuais", _str(ppc.get("vagas_anuais"))),
        ("Vagas por Turno", _str(ppc.get("vagas_turno"))),
        ("Regime de Matrícula", ppc.get("regime_matricula", "")),
        ("Periodicidade Letiva", ppc.get("periodicidade_letiva", "")),
        ("Início do Curso", ppc.get("inicio_curso", "")),
        ("Forma de Ingresso", ppc.get("formas_acesso", "")),
        ("Pré-Requisito de Ingresso", ppc.get("pre_requisito_ingresso", "")),
        ("Status do Curso", ppc.get("status_curso", "")),
        ("Tipo de Reformulação", ppc.get("tipo_reformulacao", "")),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    for i, (campo, valor) in enumerate(fields):
        _bold_cell(table.rows[i].cells[0], campo, "D9D9D9")
        table.rows[i].cells[1].text = _str(valor)
    return table


def _build_tabela_indicadores(doc: Any, ppc: dict) -> Any:
    """Tabela de indicadores de qualidade: CC / CPC / ENADE / IGC."""
    table = doc.add_table(rows=2, cols=4)
    headers = ["CC", "CPC", "ENADE", "IGC"]
    values = [
        _str(ppc.get("conceito_cc")),
        _str(ppc.get("conceito_cpc")),
        _str(ppc.get("conceito_enade")),
        _str(ppc.get("igc")),
    ]
    for i, h in enumerate(headers):
        _bold_cell(table.rows[0].cells[i], h, "D9D9D9")
        table.rows[1].cells[i].text = values[i]
    return table


def _build_tabela_nucleo(doc: Any, componentes: list[dict]) -> Any:
    """Tabela de divisão da carga horária por núcleo curricular."""
    totais: dict[str, dict] = defaultdict(lambda: {"ch_relogio": 0, "ch_aula": 0, "count": 0})
    total_relogio = 0
    total_aula = 0

    for c in componentes:
        nucleo = c.get("nucleo_curricular") or "Sem Núcleo"
        ch_r = c.get("ch_total_relogio") or 0
        ch_a = c.get("ch_total_aula") or 0
        totais[nucleo]["ch_relogio"] += ch_r
        totais[nucleo]["ch_aula"] += ch_a
        totais[nucleo]["count"] += 1
        total_relogio += ch_r
        total_aula += ch_a

    nucleos = list(totais.items())
    table = doc.add_table(rows=len(nucleos) + 2, cols=4)

    # Cabeçalho
    headers = ["Núcleo Curricular", "CH (h/r)", "CH (h/a)", "Nº Componentes"]
    for i, h in enumerate(headers):
        _bold_cell(table.rows[0].cells[i], h, "1F4E79")
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, (nucleo, dados) in enumerate(nucleos, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = nucleo
        row.cells[1].text = _str(dados["ch_relogio"])
        row.cells[2].text = _str(dados["ch_aula"])
        row.cells[3].text = _str(dados["count"])

    # Totalizador
    last_row = table.rows[-1]
    _bold_cell(last_row.cells[0], "TOTAL", "D9D9D9")
    last_row.cells[1].text = _str(total_relogio)
    last_row.cells[2].text = _str(total_aula)
    last_row.cells[3].text = _str(len(componentes))

    return table


def _build_tabela_matriz(doc: Any, componentes: list[dict]) -> Any:
    """
    Matriz curricular: agrupa componentes por período e exibe uma sub-tabela
    com cabeçalho de período para cada grupo.
    Gera uma única tabela contínua com linhas separadoras de período.
    """
    por_periodo: dict[int, list[dict]] = defaultdict(list)
    for c in componentes:
        por_periodo[c.get("periodo") or 0].append(c)

    headers = ["Componente Curricular", "CH (h/r)", "CH (h/a)", "Créd.", "Pré-req.", "Co-req."]
    total_rows = 2 + sum(1 + len(comps) for comps in por_periodo.values())  # cabeçalho + período + comps
    table = doc.add_table(rows=1, cols=6)

    # Cabeçalho global
    for i, h in enumerate(headers):
        _bold_cell(table.rows[0].cells[i], h, "1F4E79")
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for periodo in sorted(por_periodo.keys()):
        comps = por_periodo[periodo]

        # Linha de separador de período
        sep_row = table.add_row()
        sep_cell = sep_row.cells[0]
        _bold_cell(sep_cell, f"{'Período' if periodo else 'Sem Período'} {periodo if periodo else ''}", "BDD7EE")
        sep_cell.merge(sep_row.cells[-1])

        for c in comps:
            data_row = table.add_row()
            data_row.cells[0].text = c.get("nome") or ""
            data_row.cells[1].text = _str(c.get("ch_total_relogio"))
            data_row.cells[2].text = _str(c.get("ch_total_aula"))
            data_row.cells[3].text = _str(c.get("creditos"))
            data_row.cells[4].text = c.get("pre_requisito_codigo") or "—"
            data_row.cells[5].text = c.get("co_requisito_codigo") or "—"

    return table


def _build_tabela_integralizacao(doc: Any, ppc: dict, componentes: list[dict]) -> Any:
    """Tabela resumo de integralização curricular."""
    obrig = [c for c in componentes if c.get("tipo", "").lower() not in ("optativa",)]
    optativas = [c for c in componentes if c.get("tipo", "").lower() == "optativa"]
    extensao = [c for c in componentes if (c.get("ch_extensao") or 0) > 0]

    def soma_relogio(lst: list[dict]) -> int:
        return sum(c.get("ch_total_relogio") or 0 for c in lst)

    fields = [
        ("Componentes Obrigatórios (h/r)", soma_relogio(obrig), len(obrig)),
        ("Componentes Optativos (h/r)", soma_relogio(optativas), len(optativas)),
        ("CH de Extensão (h/r)", _str(ppc.get("ch_extensao")), "—"),
        ("Atividades Complementares (h/r)", _str(ppc.get("atividades_complementares")), "—"),
        ("Total Geral (h/r)", _str(ppc.get("ch_total_relogio")), len(componentes)),
    ]

    table = doc.add_table(rows=len(fields) + 1, cols=3)

    # Cabeçalho
    for i, h in enumerate(["Discriminação", "Carga Horária (h/r)", "Nº Componentes"]):
        _bold_cell(table.rows[0].cells[i], h, "D9D9D9")

    for row_idx, (disc, ch, count) in enumerate(fields, start=1):
        table.rows[row_idx].cells[0].text = disc
        table.rows[row_idx].cells[1].text = _str(ch)
        table.rows[row_idx].cells[2].text = _str(count)

    return table


def _build_tabela_optativas(doc: Any, componentes: list[dict]) -> Any:
    """Tabela de componentes optativos."""
    optativas = [c for c in componentes if c.get("tipo", "").lower() == "optativa"]
    table = doc.add_table(rows=len(optativas) + 1, cols=4)

    for i, h in enumerate(["Componente Curricular", "CH (h/r)", "CH (h/a)", "Créditos"]):
        _bold_cell(table.rows[0].cells[i], h, "D9D9D9")

    for row_idx, c in enumerate(optativas, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = c.get("nome") or ""
        row.cells[1].text = _str(c.get("ch_total_relogio"))
        row.cells[2].text = _str(c.get("ch_total_aula"))
        row.cells[3].text = _str(c.get("creditos"))

    return table


def _build_tabela_requisitos(doc: Any, componentes: list[dict]) -> Any:
    """Tabela de pré-requisitos e co-requisitos."""
    com_dep = [c for c in componentes
               if c.get("pre_requisito_codigo") or c.get("co_requisito_codigo")]
    table = doc.add_table(rows=len(com_dep) + 1, cols=3)

    for i, h in enumerate(["Componente Curricular", "Pré-Requisito", "Co-Requisito"]):
        _bold_cell(table.rows[0].cells[i], h, "D9D9D9")

    for row_idx, c in enumerate(com_dep, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = c.get("nome") or ""
        row.cells[1].text = c.get("pre_requisito_codigo") or "—"
        row.cells[2].text = c.get("co_requisito_codigo") or "—"

    return table


def _build_tabela_ementa_componente(doc: Any, componente: dict) -> Any:
    """
    Gera uma tabela de ementa para um único componente curricular.
    Formato baseado no modelo preenchido (doc_ppc_modelo_preenchido.docx).

    Estrutura:
      L1: Componente Curricular: <nome>     | Créditos: <N>
      L2: Pré-requisitos: <valor>           | (mesclada)
      L3: CH Total: <N>  Teórica: <N>  Prática: <N>  Extensão: <N>
      L4: Ementa: <texto>                   | (mesclada)
      L5: Referências Básicas: <texto>      | (mesclada)
      L6: Referências Complementares: <tex> | (mesclada)
    """
    bibs_basicas = [
        b["referencia_texto"] for b in componente.get("bibliografias", [])
        if b.get("tipo", "").lower() == "básica"
    ]
    bibs_comp = [
        b["referencia_texto"] for b in componente.get("bibliografias", [])
        if b.get("tipo", "").lower() == "complementar"
    ]

    ch_txt = (
        f"Carga horária (h/r): Total ({componente.get('ch_total_relogio') or 0}) "
        f"Teórica ({componente.get('ch_teorica') or 0}) "
        f"Prática ({componente.get('ch_pratica') or 0}) "
        f"Extensão ({componente.get('ch_extensao') or 0})"
    )
    pre_req = componente.get("pre_requisito_codigo") or "—"

    table = doc.add_table(rows=6, cols=2)

    # L1: nome e créditos
    table.rows[0].cells[0].text = f"Componente Curricular: {componente.get('nome', '')}"
    table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[0].cells[1].text = f"Créditos: {componente.get('creditos', '')}"
    table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    # L2: pré-requisitos (mesclada)
    table.rows[1].cells[0].merge(table.rows[1].cells[1])
    table.rows[1].cells[0].text = f"Pré-requisitos: {pre_req}"

    # L3: carga horária (mesclada)
    table.rows[2].cells[0].merge(table.rows[2].cells[1])
    table.rows[2].cells[0].text = ch_txt

    # L4: ementa (mesclada)
    table.rows[3].cells[0].merge(table.rows[3].cells[1])
    table.rows[3].cells[0].text = f"Ementa: {componente.get('ementa', '')}"

    # L5: referências básicas (mesclada)
    table.rows[4].cells[0].merge(table.rows[4].cells[1])
    refs_basicas_txt = "Referências Básicas: " + "  ".join(bibs_basicas) if bibs_basicas else "Referências Básicas: —"
    table.rows[4].cells[0].text = refs_basicas_txt

    # L6: referências complementares (mesclada)
    table.rows[5].cells[0].merge(table.rows[5].cells[1])
    refs_comp_txt = "Referências Complementares: " + "  ".join(bibs_comp) if bibs_comp else "Referências Complementares: —"
    table.rows[5].cells[0].text = refs_comp_txt

    return table


def _build_tabela_coordenacao(doc: Any, coordenacao: dict) -> Any:
    """Tabela de perfil do coordenador do curso."""
    fields = [
        ("Nome do Professor", coordenacao.get("nome_professor", "")),
        ("Regime de Trabalho", coordenacao.get("regime_trabalho", "")),
        ("CH Semanal de Coordenação", _str(coordenacao.get("ch_semanal_coordenacao"))),
        ("Tempo de Exercício na IES", coordenacao.get("tempo_exercicio_ies", "")),
        ("Tempo como Coordenador", coordenacao.get("tempo_coordenacao_curso", "")),
        ("Qualificação", coordenacao.get("qualificacao", "")),
        ("Titulação", coordenacao.get("titulacao", "")),
        ("Grupos de Pesquisa", coordenacao.get("grupos_pesquisa", "")),
        ("Linhas de Pesquisa", coordenacao.get("linhas_pesquisa", "")),
        ("Experiência Profissional (anos)", _str(coordenacao.get("experiencia_profissional"))),
        ("Experiência em Gestão", coordenacao.get("experiencia_gestao", "")),
        ("E-mail", coordenacao.get("email", "")),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    for i, (campo, valor) in enumerate(fields):
        _bold_cell(table.rows[i].cells[0], campo, "D9D9D9")
        table.rows[i].cells[1].text = _str(valor)
    return table


def _build_tabela_docentes(doc: Any, docentes: list[dict]) -> Any:
    """Tabela do corpo docente."""
    table = doc.add_table(rows=len(docentes) + 1, cols=4)

    for i, h in enumerate(["Nome", "Titulação", "Regime de Trabalho", "Componentes Ministrados"]):
        _bold_cell(table.rows[0].cells[i], h, "1F4E79")
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, d in enumerate(docentes, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = d.get("nome") or ""
        row.cells[1].text = d.get("titulacao") or ""
        row.cells[2].text = d.get("regime_trabalho") or ""
        row.cells[3].text = ", ".join(d.get("componentes_ministrados") or []) or "—"

    return table


def _build_tabela_colegiado_nde(doc: Any, membros: list[dict], tipo: str) -> Any:
    """
    Tabela de composição do Colegiado ou NDE.
    Filtra membros pelo tipo e gera tabela: Nome / Cargo / Portaria.
    """
    filtrados = [m for m in membros if m.get("tipo", "").lower() == tipo.lower()]
    table = doc.add_table(rows=len(filtrados) + 1, cols=3)

    for i, h in enumerate(["Nome", "Cargo", "Portaria"]):
        _bold_cell(table.rows[0].cells[i], h, "D9D9D9")

    for row_idx, m in enumerate(filtrados, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = m.get("nome") or ""
        row.cells[1].text = m.get("cargo") or ""
        row.cells[2].text = m.get("portaria") or "—"

    return table


def _build_placeholder_fluxograma(doc: Any, label: str) -> Any:
    """
    Insere um placeholder textual para fluxogramas (geração futura via imagem).
    Cria uma tabela de 1 célula como "caixa" de aviso.
    """
    table = doc.add_table(rows=1, cols=1)
    _set_cell_bg(table.rows[0].cells[0], "FFF2CC")
    table.rows[0].cells[0].text = f"[FLUXOGRAMA — {label} — A ser gerado automaticamente]"
    return table


# ── Dispatcher central ───────────────────────────────────────────────────────

def _replace_table_placeholders(doc: Any, context: dict) -> None:
    """
    Varre todos os parágrafos do documento e substitui cada tag {tabela_*}
    pela tabela correspondente gerada com python-docx.

    Args:
        doc:     Documento python-docx já renderizado pelo docxtpl.
        context: Contexto com as coleções de dados (_ppc, _componentes, etc.).
    """
    ppc = context.get("_ppc") or {}
    coordenacao = context.get("_coordenacao") or {}
    membros = context.get("_membros") or []
    componentes = context.get("_componentes") or []
    docentes = context.get("_docentes") or []

    # Itera sobre uma cópia da lista para suportar modificação durante iteração
    for para in list(doc.paragraphs):
        tag = para.text.strip().lstrip("{").rstrip("}").strip()

        if tag == "tabela_isntituicao":
            _insert_table_after_paragraph(_build_tabela_instituicao(doc, ppc), para)
            _remove_paragraph(para)

        elif tag == "tabela_curso":
            _insert_table_after_paragraph(_build_tabela_curso(doc, ppc), para)
            _remove_paragraph(para)

        elif tag == "tabela_indicadores_qualidade_curso":
            _insert_table_after_paragraph(_build_tabela_indicadores(doc, ppc), para)
            _remove_paragraph(para)

        elif tag == "tabela_divisao_componentes_por_nucleo":
            _insert_table_after_paragraph(_build_tabela_nucleo(doc, componentes), para)
            _remove_paragraph(para)

        elif tag == "tabela_componentes_por_periodo":
            _insert_table_after_paragraph(_build_tabela_matriz(doc, componentes), para)
            _remove_paragraph(para)

        elif tag == "tabela_integralizacao":
            _insert_table_after_paragraph(_build_tabela_integralizacao(doc, ppc, componentes), para)
            _remove_paragraph(para)

        elif tag == "tabela_optativas":
            _insert_table_after_paragraph(_build_tabela_optativas(doc, componentes), para)
            _remove_paragraph(para)

        elif tag == "tabela_requisitos":
            _insert_table_after_paragraph(_build_tabela_requisitos(doc, componentes), para)
            _remove_paragraph(para)

        elif tag == "tabela_ementa":
            # Gera uma tabela por componente, inserindo-as na ordem de período
            sorted_comps = sorted(componentes, key=lambda c: (c.get("periodo") or 0, c.get("nome") or ""))
            # Insere da última para a primeira para preservar a ordem após addnext()
            for comp in reversed(sorted_comps):
                tabela_ementa = _build_tabela_ementa_componente(doc, comp)
                _insert_table_after_paragraph(tabela_ementa, para)
            _remove_paragraph(para)

        elif tag == "tabela_coordenacao":
            _insert_table_after_paragraph(_build_tabela_coordenacao(doc, coordenacao), para)
            _remove_paragraph(para)

        elif tag == "tabela_docentes":
            _insert_table_after_paragraph(_build_tabela_docentes(doc, docentes), para)
            _remove_paragraph(para)

        elif tag == "tabela_composicao_colegiado":
            _insert_table_after_paragraph(_build_tabela_colegiado_nde(doc, membros, "Colegiado"), para)
            _remove_paragraph(para)

        elif tag == "tabela_composicao_nde":
            _insert_table_after_paragraph(_build_tabela_colegiado_nde(doc, membros, "NDE"), para)
            _remove_paragraph(para)

        elif tag == "fluxograma_periodos":
            _insert_table_after_paragraph(_build_placeholder_fluxograma(doc, "Matriz por Períodos"), para)
            _remove_paragraph(para)

        elif tag == "fluxograma_relacoes":
            _insert_table_after_paragraph(_build_placeholder_fluxograma(doc, "Dependências entre Componentes"), para)
            _remove_paragraph(para)

        else:
            # Loga variáveis que ficaram sem substituição (indicativo de template desatualizado)
            if para.text.strip().startswith("{") and para.text.strip().endswith("}"):
                logger.warning("Tag não reconhecida no template: %s", para.text.strip())


# ─────────────────────────────────────────────────────────────────────────────
# 5. Orquestrador principal
# ─────────────────────────────────────────────────────────────────────────────

def generate_document(ppc_id: str) -> Path:
    """
    Gera o documento DOCX completo para o PPC informado.

    Fluxo:
      1. Busca dados via carregar_ppc()
      2. Constrói o contexto de variáveis
      3. Prepara o template Jinja2 (conversão {var} → {{ var }})
      4. Renderiza variáveis simples com docxtpl
      5. Substitui placeholders de tabelas via python-docx
      6. Salva o documento final em exports/

    Args:
        ppc_id: UUID do PPC a ser exportado.

    Returns:
        Caminho absoluto do arquivo DOCX gerado.

    Raises:
        DocumentGenerationError: Em caso de falha em qualquer etapa.
    """
    if not TEMPLATE_PATH.exists():
        raise DocumentGenerationError(f"Template não encontrado: {TEMPLATE_PATH}")

    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

    logger.info("Iniciando geração de documento para PPC %s", ppc_id)

    # 1. Busca dados
    payload = fetch_ppc_data(ppc_id)
    nome_curso = (payload.get("ppc") or {}).get("nome_curso", "PPC")

    # 2. Constrói contexto
    context = build_context(payload)

    # Usa diretório temporário para arquivos intermediários
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        jinja_tpl_path = tmp / "template_jinja.docx"
        intermediate_path = tmp / "intermediate.docx"

        # 3. Converte template original → Jinja2
        try:
            _prepare_jinja_template(TEMPLATE_PATH, jinja_tpl_path)
        except Exception as exc:
            raise DocumentGenerationError(f"Falha ao preparar template Jinja2: {exc}") from exc

        # 4. Renderiza variáveis de texto com docxtpl
        try:
            render_template(jinja_tpl_path, context, intermediate_path)
        except Exception as exc:
            raise DocumentGenerationError(f"Falha ao renderizar variáveis de texto: {exc}") from exc

        # 5. Substitui placeholders de tabelas via python-docx
        try:
            doc = Document(str(intermediate_path))
            _replace_table_placeholders(doc, context)
        except Exception as exc:
            raise DocumentGenerationError(f"Falha ao gerar tabelas: {exc}") from exc

        # 6. Salva documento final
        safe_name = re.sub(r"[^\w\-]", "_", nome_curso)[:60]
        output_filename = f"PPC_{safe_name}_{ppc_id[:8]}.docx"
        output_path = EXPORTS_DIR / output_filename

        try:
            doc.save(str(output_path))
        except Exception as exc:
            raise DocumentGenerationError(f"Falha ao salvar documento final: {exc}") from exc

    logger.info("Documento gerado com sucesso: %s", output_path)
    return output_path
